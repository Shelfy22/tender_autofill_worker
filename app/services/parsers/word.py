from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document

from app.config import Settings


def _convert(path: Path, target_extension: str, settings: Settings) -> Path:
    output_dir = path.parent / "converted"
    output_dir.mkdir(exist_ok=True)
    subprocess.run(
        [
            settings.libreoffice_binary,
            "--headless",
            "--convert-to",
            target_extension,
            "--outdir",
            str(output_dir),
            str(path),
        ],
        capture_output=True,
        timeout=settings.conversion_timeout_seconds,
        check=True,
    )
    converted = output_dir / f"{path.stem}.{target_extension}"
    if not converted.exists():
        raise RuntimeError(f"LibreOffice не создал {converted.name}")
    return converted


def extract_word_text(path: Path, file_type: str, settings: Settings) -> tuple[str, str, list[str]]:
    warnings: list[str] = []
    docx_path = path
    if file_type == "doc":
        docx_path = _convert(path, "docx", settings)
        warnings.append("Старый DOC преобразован LibreOffice в DOCX.")
    document = Document(docx_path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            if any(values):
                parts.append(" | ".join(values))
    text = "\n".join(parts).strip()
    quality = len(text) >= 500 and any(char.isalnum() for char in text)
    if not quality:
        warnings.append(f"Word не дал полезный текст: длина {len(text)}.")
    return (text if quality else ""), ("ok" if quality else "word_quality_failed"), warnings
